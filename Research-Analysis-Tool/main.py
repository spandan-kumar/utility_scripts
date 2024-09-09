import openai
import os
import docx
from PyPDF2 import PdfReader

openai.api_key = 'key'

def extract_text_from_pdf(pdf_path):
    try:
        reader = PdfReader(pdf_path)
        text = ''
        for page in reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        print(f"Error reading {pdf_path}: {e}")
        return None

def analyze_text(text):
    try:
        prompt = f"Extract the paper title, journal details, year of publishing, tools, framework, algorithm, methodology, inferences, and research gaps from the following research paper and nothing else:\n\n{text}"
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an expert in research analysis and give concise, to-the-point answers."},
                {"role": "user", "content": prompt}
            ]
        )
        return response['choices'][0]['message']['content']
    except Exception as e:
        print(f"Error analyzing text: {e}")
        return "Analysis failed due to an error."

def save_to_docx(analysis, output_path):
    try:
        doc = docx.Document()
        doc.add_heading('Research Paper Analysis', 0)
        doc.add_paragraph(analysis)
        doc.save(output_path)
    except Exception as e:
        print(f"Error saving DOCX file: {e}")

folder_path = "folder_contaning_papers"
output_file = ".docs_path"

all_analyses = ""
for filename in os.listdir(folder_path):
    file_path = os.path.join(folder_path, filename)
    if filename.endswith('.pdf'):
        text = extract_text_from_pdf(file_path)
        if text:
            analysis = analyze_text(text)
            all_analyses += f"Analysis for {filename}:\n\n{analysis}\n\n{'-'*40}\n\n"

save_to_docx(all_analyses, output_file)

print(f"Analysis saved to {output_file}")
